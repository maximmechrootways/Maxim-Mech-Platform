"""
GX10 local RAG service.

- Sweeps DATA_DIR/inbox (recursively) for files and folders dropped via the SMB share
  or copied from USB. Folder structure is preserved: the top-level folder becomes the
  "project", nested folders become the folder path. Files move to DATA_DIR/library,
  get extracted, chunked, embedded via Ollama, and stored in pgvector.
- Unsupported file types are still archived (status 'stored') so whole project folders
  can be dropped as-is; they're downloadable but not semantically searchable.
- Exposes an authenticated HTTP API consumed by the Maxim backend (Frank):
    GET  /health
    POST /search                {"query": str, "limit": int, "project": str?}
    GET  /documents             list ingested documents
    GET  /documents/{id}/file   raw file bytes (for preview/download, proxied by Maxim backend)
    POST /documents/{id}/reindex
    GET  /projects              per-project document counts/sizes
    POST /upload                multipart upload (drops files into the inbox pipeline)
- Serves a browser drop-box GUI at / (static/index.html) so anyone on the LAN can
  plug in a USB, open the page, and drag whole project folders in.
"""

import hashlib
import logging
import mimetypes
import os
import queue
import re
import shutil
import threading
import time
import uuid
from contextlib import asynccontextmanager
from pathlib import Path

import fitz  # PyMuPDF
import httpx
from docx import Document as DocxDocument
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
from pgvector.psycopg import register_vector
from psycopg_pool import ConnectionPool
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("local-rag")

DATABASE_URL = os.environ["DATABASE_URL"]
OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://ollama:11434")
EMBED_MODEL = os.environ.get("EMBED_MODEL", "bge-m3")
EMBED_DIMENSIONS = int(os.environ.get("EMBED_DIMENSIONS", "1024"))
API_KEY = os.environ["API_KEY"]
# Short human-friendly code that gates the drag-and-drop upload page (entered once per browser).
UPLOAD_PASSCODE = os.environ.get("UPLOAD_PASSCODE", "")
DATA_DIR = Path(os.environ.get("DATA_DIR", "/data"))
INBOX_DIR = DATA_DIR / "inbox"
LIBRARY_DIR = DATA_DIR / "library"
# Uploads are staged here first, then moved into the inbox in one rename so the
# sweep loop never sees a half-written file.
UPLOAD_TMP_DIR = DATA_DIR / "upload-tmp"
STATIC_DIR = Path(__file__).parent / "static"
SWEEP_INTERVAL_SECONDS = int(os.environ.get("SWEEP_INTERVAL_SECONDS", "10"))

# Mirrors the cloud pipeline in documentIngestionService.ts
WORDS_PER_CHUNK = 400
OVERLAP_WORDS = 50

# Indexable (searchable) types. Everything else is archived as-is.
INDEXABLE_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
# OS/copy junk that should be silently deleted so folders can empty out.
JUNK_FILE_NAMES = {"thumbs.db", "desktop.ini", "autorun.inf", ".ds_store"}
JUNK_DIR_NAMES = {"system volume information", "$recycle.bin", ".trashes", ".spotlight-v100"}

pool = ConnectionPool(DATABASE_URL, min_size=1, max_size=8, open=False, configure=register_vector)

shutdown_event = threading.Event()
# Embed off the inbox sweeper thread so a big PDF can't freeze 200+ waiting files.
_embed_queue: queue.Queue[tuple] = queue.Queue()
_embed_workers_started = False


def _ensure_embed_workers(n: int = 1) -> None:
    """Start background embed workers once (single worker keeps Ollama from thrashing)."""
    global _embed_workers_started
    if _embed_workers_started:
        return
    _embed_workers_started = True

    def worker() -> None:
        while not shutdown_event.is_set():
            try:
                item = _embed_queue.get(timeout=1.0)
            except queue.Empty:
                continue
            try:
                doc_id, file_path, name = item
                embed_document(doc_id, Path(file_path), name)
            except Exception:
                log.exception("Embed worker crashed on %s", item)
            finally:
                _embed_queue.task_done()

    for i in range(max(1, n)):
        threading.Thread(target=worker, name=f"embed-worker-{i}", daemon=True).start()
    log.info("Started %d background embed worker(s)", max(1, n))


def enqueue_embed(doc_id, file_path: Path, name: str) -> None:
    _ensure_embed_workers()
    _embed_queue.put((doc_id, str(file_path), name))
    log.info("Queued embed for %s (%d waiting in embed queue)", name, _embed_queue.qsize())


# ---------------------------------------------------------------- embedding

def ensure_model_pulled() -> None:
    """Pull the embedding model on first boot so ingestion never fails on a cold start."""
    with httpx.Client(base_url=OLLAMA_URL, timeout=None) as client:
        tags = client.get("/api/tags").json()
        names = {m["name"].split(":")[0] for m in tags.get("models", [])}
        if EMBED_MODEL.split(":")[0] not in names:
            log.info("Pulling embedding model %s (first run, may take a few minutes)", EMBED_MODEL)
            client.post("/api/pull", json={"model": EMBED_MODEL})


def embed_texts(texts: list[str]) -> list[list[float]]:
    with httpx.Client(base_url=OLLAMA_URL, timeout=300) as client:
        res = client.post("/api/embed", json={"model": EMBED_MODEL, "input": texts})
        res.raise_for_status()
        embeddings = res.json()["embeddings"]
    for e in embeddings:
        if len(e) != EMBED_DIMENSIONS:
            raise RuntimeError(
                f"Model returned {len(e)}-dim vectors but EMBED_DIMENSIONS={EMBED_DIMENSIONS}; "
                "update .env and schema.sql to match the model."
            )
    return embeddings


# ---------------------------------------------------------------- extraction

def chunk_words(text: str) -> list[str]:
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i : i + WORDS_PER_CHUNK])
        if len(chunk.strip()) > 30:
            chunks.append(chunk)
        if i + WORDS_PER_CHUNK >= len(words):
            break
        i += WORDS_PER_CHUNK - OVERLAP_WORDS
    return chunks


def metadata_search_text(name: str, project: str = "", folder_path: str = "") -> str:
    """Build a searchable blurb from filename/path so drawings without OCR still match."""
    stem = Path(name).stem
    human = re.sub(r"[_\-.]+", " ", stem)
    human = re.sub(r"\s+", " ", human).strip()
    bits = [
        f"File name: {name}",
        f"Title keywords: {human}" if human else "",
        f"Project: {project}" if project else "",
        f"Folder path: {folder_path}" if folder_path else "",
        "Local archive document. Drawing sheet. Plan. Spec. Shop drawing.",
    ]
    # Help queries like "electrical drawing" hit names that contain those words
    lower = f"{name} {human} {folder_path} {project}".lower()
    for hint in (
        "electrical",
        "mechanical",
        "plumbing",
        "hvac",
        "drawing",
        "dwg",
        "schematic",
        "single line",
        "panel",
        "shop drawing",
        "submittal",
    ):
        if hint in lower and hint not in " ".join(bits).lower():
            bits.append(f"Tag: {hint}")
    return "\n".join(b for b in bits if b)


def extract_pages(path: Path) -> list[tuple[int, str]]:
    """Returns (page_number, text) pairs. Non-paged formats are one 'page'."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        with fitz.open(path) as doc:
            return [(i + 1, page.get_text()) for i, page in enumerate(doc) if page.get_text().strip()]
    if ext == ".docx":
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return [(1, text)] if text.strip() else []
    return [(1, path.read_text(errors="replace"))]


# ---------------------------------------------------------------- ingestion

def _mark_stored(doc_id, note: str) -> None:
    """Keep the file archived/downloadable, but not semantically searchable."""
    with pool.connection() as conn:
        conn.execute(
            """UPDATE local_documents
               SET status = 'stored', chunk_count = 0, error = %s, updated_at = now()
               WHERE id = %s""",
            (note[:2000], doc_id),
        )


def embed_document(doc_id, file_path: Path, name: str) -> None:
    """Extract, chunk, embed and store chunks for an existing local_documents row.

    Always indexes filename/project/folder so electrical drawings (and other
    image-only PDFs) remain findable by name even with no text layer.
    """
    try:
        with pool.connection() as conn:
            meta_row = conn.execute(
                "SELECT project, folder_path FROM local_documents WHERE id = %s",
                (doc_id,),
            ).fetchone()
        project = (meta_row[0] if meta_row else "") or ""
        folder_path = (meta_row[1] if meta_row else "") or ""

        all_chunks: list[tuple[int, int, str]] = []  # (page, chunk_index, content)
        # page 0 = metadata chunk (filename / path)
        all_chunks.append((0, 0, metadata_search_text(name, project, folder_path)))
        idx = 1

        pages = extract_pages(file_path)
        body_chunks = 0
        for page_num, text in pages:
            for content in chunk_words(text):
                all_chunks.append((page_num, idx, content))
                idx += 1
                body_chunks += 1

        filename_only = body_chunks == 0

        BATCH = 64
        with pool.connection() as conn:
            conn.execute("DELETE FROM local_chunks WHERE document_id = %s", (doc_id,))
            for i in range(0, len(all_chunks), BATCH):
                batch = all_chunks[i : i + BATCH]
                embeddings = embed_texts([c[2] for c in batch])
                with conn.cursor() as cur:
                    cur.executemany(
                        """INSERT INTO local_chunks (document_id, content, page_number, chunk_index, embedding)
                           VALUES (%s, %s, %s, %s, %s)""",
                        [(doc_id, c[2], c[0], c[1], embeddings[j]) for j, c in enumerate(batch)],
                    )
                log.info(
                    "  %s: embedded chunks %d-%d/%d",
                    name,
                    i + 1,
                    min(i + BATCH, len(all_chunks)),
                    len(all_chunks),
                )
            note = (
                "Indexed by filename/path only (no PDF text layer)."
                if filename_only
                else None
            )
            conn.execute(
                """UPDATE local_documents
                   SET status = 'ingested', chunk_count = %s, error = %s, updated_at = now()
                   WHERE id = %s""",
                (len(all_chunks), note, doc_id),
            )
        log.info(
            "Ingested %s (%d chunks%s)",
            name,
            len(all_chunks),
            ", filename-only" if filename_only else "",
        )
    except Exception as exc:
        log.exception("Ingestion failed for %s", name)
        with pool.connection() as conn:
            conn.execute(
                "UPDATE local_documents SET status = 'failed', error = %s, updated_at = now() WHERE id = %s",
                (str(exc)[:2000], doc_id),
            )


def process_inbox_file(path: Path) -> None:
    """Move one inbox file into the library (preserving folder structure) and index it."""
    rel = path.relative_to(INBOX_DIR)
    # Top-level folder = project; nested path = folder within the project.
    project = rel.parts[0] if len(rel.parts) > 1 else ""
    folder_path = "/".join(rel.parts[1:-1]) if len(rel.parts) > 2 else ""
    name = path.name
    size = path.stat().st_size
    if size == 0:
        log.warning("Removing empty inbox file %s (bad upload)", rel.as_posix())
        path.unlink(missing_ok=True)
        return
    sha = hashlib.sha256(path.read_bytes()).hexdigest()
    indexable = path.suffix.lower() in INDEXABLE_EXTENSIONS

    with pool.connection() as conn:
        dup = conn.execute(
            "SELECT id FROM local_documents WHERE sha256 = %s AND project = %s AND status IN ('ingested', 'stored')",
            (sha, project),
        ).fetchone()
        if dup:
            log.info("Skipping duplicate %s in project '%s' (already stored as %s)", name, project, dup[0])
            path.unlink(missing_ok=True)
            return

    dest = LIBRARY_DIR / rel
    dest.parent.mkdir(parents=True, exist_ok=True)
    if dest.exists():
        dest = dest.with_name(f"{uuid.uuid4().hex[:8]}-{dest.name}")
    shutil.move(str(path), str(dest))

    content_type = mimetypes.guess_type(name)[0] or "application/octet-stream"
    with pool.connection() as conn:
        row = conn.execute(
            """INSERT INTO local_documents (name, file_path, content_type, size_bytes, sha256, status, project, folder_path)
               VALUES (%s, %s, %s, %s, %s, 'pending', %s, %s)
               ON CONFLICT (file_path) DO UPDATE SET updated_at = now()
               RETURNING id""",
            (name, dest.relative_to(DATA_DIR).as_posix(), content_type, size, sha, project, folder_path),
        ).fetchone()
        doc_id = row[0]

    if indexable:
        # Don't block the inbox sweeper on Ollama — drain inbox first, embed in background.
        enqueue_embed(doc_id, dest, name)
    else:
        with pool.connection() as conn:
            conn.execute(
                "UPDATE local_documents SET status = 'stored', updated_at = now() WHERE id = %s",
                (doc_id,),
            )
        log.info("Archived %s (type %s not searchable, stored only)", name, path.suffix.lower() or "unknown")


# ---------------------------------------------------------------- inbox sweep

def wait_until_stable(path: Path, checks: int = 2, interval: float = 1.0, max_wait: int = 120) -> bool:
    """Don't ingest files still being copied over SMB/USB.

    Zero-byte files are treated as stable after `checks` so a bad upload cannot
    block the inbox forever (previously size>0 was required, so empties stuck).
    """
    last = -1
    stable = 0
    for _ in range(max_wait):
        if not path.exists():
            return False
        try:
            size = path.stat().st_size
        except OSError:
            return False
        if size == last:
            stable += 1
            if stable >= checks:
                return True
        else:
            stable = 0
            last = size
        time.sleep(interval)
    return False


def is_junk(path: Path) -> bool:
    if path.name.startswith((".", "~$")) or path.name.lower() in JUNK_FILE_NAMES:
        return True
    return any(part.lower() in JUNK_DIR_NAMES for part in path.relative_to(INBOX_DIR).parts[:-1])


def cleanup_empty_dirs() -> None:
    """Remove empty inbox subdirectories that haven't been touched for a couple of minutes."""
    cutoff = time.time() - 120
    dirs = sorted((p for p in INBOX_DIR.rglob("*") if p.is_dir()), key=lambda p: len(p.parts), reverse=True)
    for d in dirs:
        try:
            if not any(d.iterdir()) and d.stat().st_mtime < cutoff:
                d.rmdir()
        except OSError:
            pass


def sweep_once() -> None:
    for path in sorted(INBOX_DIR.rglob("*")):
        if shutdown_event.is_set():
            return
        if not path.is_file():
            continue
        if path.suffix.lower() == ".part":  # in-flight web upload, renamed when complete
            continue
        if is_junk(path):
            path.unlink(missing_ok=True)
            continue
        if wait_until_stable(path):
            try:
                process_inbox_file(path)
            except Exception:
                log.exception("Failed to process %s", path)
    cleanup_empty_dirs()


def sweep_loop() -> None:
    while not shutdown_event.is_set():
        try:
            sweep_once()
        except Exception:
            log.exception("Inbox sweep failed")
        shutdown_event.wait(SWEEP_INTERVAL_SECONDS)


# ---------------------------------------------------------------- API

def require_api_key(request: Request) -> None:
    if request.headers.get("x-api-key") != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")


def run_startup_migrations() -> None:
    """Idempotent column additions for deployments created before project support."""
    with pool.connection() as conn:
        conn.execute("ALTER TABLE local_documents ADD COLUMN IF NOT EXISTS project TEXT NOT NULL DEFAULT ''")
        conn.execute("ALTER TABLE local_documents ADD COLUMN IF NOT EXISTS folder_path TEXT NOT NULL DEFAULT ''")
        conn.execute("CREATE INDEX IF NOT EXISTS local_documents_project_idx ON local_documents (project)")
        # Older builds marked scanned PDFs as failed; keep them as archived/stored instead.
        conn.execute(
            """UPDATE local_documents
               SET status = 'stored',
                   error = 'No searchable text (scanned/image PDF). File is archived and downloadable.'
               WHERE status = 'failed'
                 AND (error ILIKE '%No readable text%' OR error ILIKE '%scanned%')"""
        )
        # Re-index stored "no text" files so filename/path search works (electrical drawings etc.).
        conn.execute(
            """UPDATE local_documents
               SET status = 'pending', error = NULL, updated_at = now()
               WHERE status = 'stored'
                 AND (
                   error ILIKE '%No searchable text%'
                   OR error ILIKE '%Text too sparse%'
                   OR error ILIKE '%scanned/image PDF%'
                 )"""
        )


@asynccontextmanager
async def lifespan(app: FastAPI):
    INBOX_DIR.mkdir(parents=True, exist_ok=True)
    LIBRARY_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_TMP_DIR.mkdir(parents=True, exist_ok=True)
    pool.open()
    run_startup_migrations()
    ensure_model_pulled()
    _ensure_embed_workers(1)
    # Re-queue anything left pending after a restart (e.g. mid-embed crash).
    with pool.connection() as conn:
        pending = conn.execute(
            """SELECT id, file_path, name FROM local_documents
               WHERE status = 'pending' ORDER BY created_at"""
        ).fetchall()
    for doc_id, file_path, name in pending:
        src = DATA_DIR / file_path
        if src.exists():
            enqueue_embed(doc_id, src, name)
        else:
            log.warning("Pending doc %s missing on disk: %s", doc_id, file_path)
    if pending:
        log.info("Re-queued %d pending document(s) for embedding", len(pending))
    threading.Thread(target=sweep_loop, daemon=True).start()
    log.info("Sweeping %s every %ds; embedding with %s (%d dims)", INBOX_DIR, SWEEP_INTERVAL_SECONDS, EMBED_MODEL, EMBED_DIMENSIONS)
    yield
    shutdown_event.set()
    pool.close()


app = FastAPI(title="GX10 Local RAG", lifespan=lifespan)


class SearchRequest(BaseModel):
    query: str
    limit: int = 5
    project: str | None = None


def require_upload_auth(request: Request) -> None:
    """Desktop app uses X-API-Key; LAN browser drop-box may use X-Upload-Code."""
    if request.headers.get("x-api-key") == API_KEY:
        return
    if UPLOAD_PASSCODE and request.headers.get("x-upload-code") == UPLOAD_PASSCODE:
        return
    if not UPLOAD_PASSCODE:
        raise HTTPException(status_code=401, detail="Invalid API key (or set UPLOAD_PASSCODE for browser uploads)")
    raise HTTPException(status_code=401, detail="Wrong upload code or API key")


_upload_page_path = Path(__file__).parent / "upload_page.html"
UPLOAD_PAGE = (
    _upload_page_path.read_text(encoding="utf-8")
    if _upload_page_path.exists()
    else "<!doctype html><html><body><h1>Maxim Local Archive</h1><p>Use the desktop app or SMB share to upload.</p></body></html>"
)


@app.get("/", response_class=HTMLResponse)
def upload_page():
    return UPLOAD_PAGE


def sanitize_relpath(relpath: str) -> Path:
    """Turn a client-supplied relative path into a safe path under the inbox."""
    parts = []
    for part in relpath.replace("\\", "/").split("/"):
        part = part.strip()
        if not part or part in (".", ".."):
            continue
        parts.append(part)
    if not parts:
        raise HTTPException(status_code=400, detail="Invalid file path")
    return Path(*parts)


@app.post("/upload", dependencies=[Depends(require_upload_auth)])
async def upload(file: UploadFile = File(...), relpath: str = Form(""), project: str = Form("")):
    """Accept a file into the inbox pipeline. Used by the Electron app (API key) and LAN page."""
    rel = sanitize_relpath(relpath or file.filename or "upload.bin")
    project_name = project.strip()
    if project_name:
        rel = Path(sanitize_relpath(project_name).as_posix()) / rel

    dest = INBOX_DIR / rel
    dest.parent.mkdir(parents=True, exist_ok=True)
    tmp = UPLOAD_TMP_DIR / f"{uuid.uuid4().hex}.part"

    size = 0
    try:
        with open(tmp, "wb") as out:
            while chunk := await file.read(1024 * 1024):
                out.write(chunk)
                size += len(chunk)
        if size == 0:
            raise HTTPException(status_code=400, detail="Empty file upload rejected")
        shutil.move(str(tmp), str(dest))
    finally:
        if tmp.exists():
            tmp.unlink(missing_ok=True)

    log.info("Upload received: %s (%d bytes)", rel.as_posix(), size)
    return {"ok": True, "path": rel.as_posix(), "sizeBytes": size}


@app.get("/upload-status", dependencies=[Depends(require_upload_auth)])
def upload_status():
    """Live feedback for upload UIs: inbox vs still-embedding."""
    pending_files = 0
    pending_projects: set[str] = set()
    pending_sample: list[dict] = []
    for p in INBOX_DIR.rglob("*"):
        if not p.is_file() or p.suffix.lower() == ".part" or is_junk(p):
            continue
        pending_files += 1
        try:
            rel = p.relative_to(INBOX_DIR)
            size = p.stat().st_size
            if len(rel.parts) > 1:
                pending_projects.add(rel.parts[0])
            if len(pending_sample) < 8:
                pending_sample.append({"path": rel.as_posix(), "sizeBytes": size})
        except (ValueError, OSError):
            pass
    with pool.connection() as conn:
        pending_docs, ingested, stored, failed = conn.execute(
            """SELECT
                 count(*) FILTER (WHERE status = 'pending'),
                 count(*) FILTER (WHERE status = 'ingested'),
                 count(*) FILTER (WHERE status = 'stored'),
                 count(*) FILTER (WHERE status = 'failed')
               FROM local_documents"""
        ).fetchone()
        rows = conn.execute(
            """SELECT name, project, status, chunk_count, error
               FROM local_documents ORDER BY updated_at DESC LIMIT 15"""
        ).fetchall()
    return {
        "pendingFiles": pending_files,
        "pendingDocuments": int(pending_docs),
        "embedQueue": _embed_queue.qsize(),
        "counts": {
            "ingested": int(ingested),
            "stored": int(stored),
            "failed": int(failed),
            "pending": int(pending_docs),
        },
        "pendingProjects": sorted(pending_projects),
        "pendingSample": pending_sample,
        "recent": [
            {"name": r[0], "project": r[1], "status": r[2], "chunkCount": r[3], "error": r[4]}
            for r in rows
        ],
    }


@app.get("/tree", dependencies=[Depends(require_api_key)])
def document_tree(project: str | None = None):
    """Nested project → folder → files tree for the desktop app and maximmech.com."""
    with pool.connection() as conn:
        if project:
            rows = conn.execute(
                """SELECT id, name, content_type, size_bytes, status, chunk_count, error, created_at, project, folder_path
                   FROM local_documents
                   WHERE status IN ('ingested', 'stored', 'failed', 'pending')
                     AND project ILIKE %s
                   ORDER BY project, folder_path, name""",
                (f"%{project}%",),
            ).fetchall()
        else:
            rows = conn.execute(
                """SELECT id, name, content_type, size_bytes, status, chunk_count, error, created_at, project, folder_path
                   FROM local_documents
                   WHERE status IN ('ingested', 'stored', 'failed', 'pending')
                   ORDER BY project, folder_path, name"""
            ).fetchall()

    # projectName -> folderPath -> files
    projects: dict[str, dict] = {}
    for r in rows:
        proj = r[8] or "(loose files)"
        folder = r[9] or ""
        entry = projects.setdefault(proj, {"name": proj, "folders": {}})
        folder_entry = entry["folders"].setdefault(folder, {"path": folder, "files": []})
        folder_entry["files"].append(
            {
                "id": str(r[0]),
                "name": r[1],
                "contentType": r[2],
                "sizeBytes": r[3],
                "status": r[4],
                "chunkCount": r[5],
                "error": r[6],
                "createdAt": r[7].isoformat() if r[7] else None,
                "project": r[8],
                "folderPath": r[9],
            }
        )

    tree = []
    for proj_name, proj_data in projects.items():
        folders = [
            {"path": path, "files": data["files"]}
            for path, data in sorted(proj_data["folders"].items(), key=lambda x: x[0])
        ]
        file_count = sum(len(f["files"]) for f in folders)
        tree.append({"name": proj_name, "fileCount": file_count, "folders": folders})
    tree.sort(key=lambda p: p["name"].lower())
    return {"projects": tree}


@app.get("/projects", dependencies=[Depends(require_api_key)])
def list_projects():
    with pool.connection() as conn:
        rows = conn.execute(
            """SELECT project, count(*), coalesce(sum(size_bytes), 0),
                      count(*) FILTER (WHERE status = 'ingested'), max(created_at)
               FROM local_documents
               WHERE status IN ('ingested', 'stored', 'failed', 'pending')
               GROUP BY project
               ORDER BY project"""
        ).fetchall()
    return {
        "projects": [
            {
                "name": r[0] or "(loose files)",
                "fileCount": r[1],
                "totalSizeBytes": int(r[2]),
                "searchableCount": r[3],
                "lastUpdated": r[4].isoformat() if r[4] else None,
            }
            for r in rows
            if r[0]  # skip empty project bucket for the picker
        ]
    }


@app.get("/health")
def health():
    with pool.connection() as conn:
        docs, chunks, projects = conn.execute(
            """SELECT (SELECT count(*) FROM local_documents WHERE status IN ('ingested', 'stored')),
                      (SELECT count(*) FROM local_chunks),
                      (SELECT count(DISTINCT project) FROM local_documents WHERE project <> '')"""
        ).fetchone()
    return {"ok": True, "documents": docs, "chunks": chunks, "projects": projects, "model": EMBED_MODEL}


@app.post("/search", dependencies=[Depends(require_api_key)])
def search(body: SearchRequest):
    """Hybrid search: vector similarity over chunks + filename/folder/project keyword hits.

    Image-only drawings often have no PDF text; filename indexing + ILIKE catches
    queries like "electrical drawing" when those words appear in the path/name.
    """
    query = body.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="query is required")
    limit = max(1, min(body.limit, 10))
    [embedding] = embed_texts([query])
    project_filter = (body.project or "").strip()
    tokens = [t for t in re.split(r"[^a-zA-Z0-9]+", query.lower()) if len(t) >= 3]

    with pool.connection() as conn:
        if project_filter:
            vec_rows = conn.execute(
                """SELECT c.content, c.page_number, c.chunk_index,
                          d.id, d.name, d.project, d.folder_path,
                          1 - (c.embedding <=> %s::vector) AS similarity
                   FROM local_chunks c
                   JOIN local_documents d ON d.id = c.document_id
                   WHERE d.project ILIKE %s
                   ORDER BY c.embedding <=> %s::vector
                   LIMIT %s""",
                (str(embedding), f"%{project_filter}%", str(embedding), limit * 2),
            ).fetchall()
        else:
            vec_rows = conn.execute(
                """SELECT c.content, c.page_number, c.chunk_index,
                          d.id, d.name, d.project, d.folder_path,
                          1 - (c.embedding <=> %s::vector) AS similarity
                   FROM local_chunks c
                   JOIN local_documents d ON d.id = c.document_id
                   ORDER BY c.embedding <=> %s::vector
                   LIMIT %s""",
                (str(embedding), str(embedding), limit * 2),
            ).fetchall()

        name_rows = []
        if tokens:
            # Require at least one query token to appear in name/folder/project
            like_clauses = []
            params: list = []
            for tok in tokens[:8]:
                like_clauses.append(
                    "(d.name ILIKE %s OR d.folder_path ILIKE %s OR d.project ILIKE %s)"
                )
                params.extend([f"%{tok}%", f"%{tok}%", f"%{tok}%"])
            where_name = " OR ".join(like_clauses)
            sql = f"""
                SELECT d.id, d.name, d.project, d.folder_path, d.status
                FROM local_documents d
                WHERE d.status IN ('ingested', 'stored')
                  AND ({where_name})
            """
            if project_filter:
                sql += " AND d.project ILIKE %s"
                params.append(f"%{project_filter}%")
            sql += " ORDER BY d.updated_at DESC LIMIT %s"
            params.append(limit * 3)
            name_rows = conn.execute(sql, params).fetchall()

    # Merge: prefer higher similarity; boost exact-ish filename matches
    by_doc: dict[str, dict] = {}
    for r in vec_rows:
        doc_id = str(r[3])
        sim = float(r[7])
        prev = by_doc.get(doc_id)
        if not prev or sim > prev["similarity"]:
            by_doc[doc_id] = {
                "content": r[0],
                "pageNumber": r[1],
                "chunkIndex": r[2],
                "documentId": doc_id,
                "documentName": r[4],
                "project": r[5],
                "folderPath": r[6],
                "similarity": sim,
            }

    q_lower = query.lower()
    for r in name_rows:
        doc_id = str(r[0])
        name, project, folder = r[1] or "", r[2] or "", r[3] or ""
        hay = f"{name} {folder} {project}".lower()
        hits = sum(1 for t in tokens if t in hay)
        boost = 0.55 + 0.08 * hits
        if any(t in name.lower() for t in tokens):
            boost += 0.15
        if q_lower in hay:
            boost = max(boost, 0.93)
        content = metadata_search_text(name, project, folder)
        content += "\n(Matched by file name / folder — PDF may have no readable text.)"
        prev = by_doc.get(doc_id)
        if not prev or boost > prev["similarity"]:
            by_doc[doc_id] = {
                "content": content,
                "pageNumber": 0,
                "chunkIndex": -1,
                "documentId": doc_id,
                "documentName": name,
                "project": project,
                "folderPath": folder,
                "similarity": min(boost, 0.99),
            }

    ranked = sorted(by_doc.values(), key=lambda x: x["similarity"], reverse=True)[:limit]
    return {"results": ranked}


@app.get("/documents", dependencies=[Depends(require_api_key)])
def list_documents(project: str | None = None):
    with pool.connection() as conn:
        if project:
            rows = conn.execute(
                """SELECT id, name, content_type, size_bytes, status, chunk_count, error, created_at, project, folder_path
                   FROM local_documents WHERE project ILIKE %s ORDER BY project, folder_path, name""",
                (f"%{project}%",),
            ).fetchall()
        else:
            rows = conn.execute(
                """SELECT id, name, content_type, size_bytes, status, chunk_count, error, created_at, project, folder_path
                   FROM local_documents ORDER BY project, folder_path, name"""
            ).fetchall()
    return {
        "documents": [
            {
                "id": str(r[0]), "name": r[1], "contentType": r[2], "sizeBytes": r[3],
                "status": r[4], "chunkCount": r[5], "error": r[6], "createdAt": r[7].isoformat(),
                "project": r[8], "folderPath": r[9],
            }
            for r in rows
        ]
    }


def _safe_data_path(rel: str) -> Path | None:
    """Resolve a DB file_path under DATA_DIR; reject path traversal."""
    try:
        path = (DATA_DIR / rel).resolve()
        data_root = DATA_DIR.resolve()
        if path != data_root and data_root not in path.parents:
            return None
        return path
    except OSError:
        return None


def _unlink_file_and_prune(rel: str) -> bool:
    path = _safe_data_path(rel)
    removed = False
    if path and path.is_file():
        try:
            path.unlink()
            removed = True
        except OSError as exc:
            log.warning("Could not delete file %s: %s", path, exc)
    # Prune empty parent dirs up to library/ (do not remove library itself)
    if path:
        parent = path.parent
        library_root = LIBRARY_DIR.resolve()
        while parent != library_root and library_root in parent.parents:
            try:
                parent.rmdir()
            except OSError:
                break
            parent = parent.parent
    return removed


@app.get("/documents/{doc_id}/file", dependencies=[Depends(require_api_key)])
def get_file(doc_id: str):
    with pool.connection() as conn:
        row = conn.execute(
            "SELECT name, file_path, content_type FROM local_documents WHERE id = %s", (doc_id,)
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    path = DATA_DIR / row[1]
    if not path.exists():
        raise HTTPException(status_code=410, detail="File missing on disk")
    return FileResponse(path, media_type=row[2] or "application/octet-stream", filename=row[0])


@app.delete("/documents/{doc_id}", dependencies=[Depends(require_api_key)])
def delete_document(doc_id: str):
    """Remove one document: disk file + DB row (chunks cascade)."""
    with pool.connection() as conn:
        row = conn.execute(
            "SELECT name, file_path, project FROM local_documents WHERE id = %s", (doc_id,)
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")
        name, file_path, project = row[0], row[1], row[2]
        conn.execute("DELETE FROM local_documents WHERE id = %s", (doc_id,))
    removed = _unlink_file_and_prune(file_path)
    log.info("Deleted document %s (%s) project=%s file_removed=%s", doc_id, name, project, removed)
    return {"deleted": True, "id": doc_id, "name": name, "project": project, "fileRemoved": removed}


@app.delete("/projects/{project_name}", dependencies=[Depends(require_api_key)])
def delete_project(project_name: str):
    """Remove every document in a project (exact name match), including disk files."""
    return _delete_project_by_name(project_name)


@app.delete("/projects", dependencies=[Depends(require_api_key)])
def delete_project_query(name: str = ""):
    """Same as path delete; query form avoids encoding issues with spaces in names."""
    return _delete_project_by_name(name)


def _delete_project_by_name(project_name: str):
    name = (project_name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="project name is required")
    with pool.connection() as conn:
        rows = conn.execute(
            "SELECT id, file_path FROM local_documents WHERE project = %s",
            (name,),
        ).fetchall()
        if not rows:
            raise HTTPException(status_code=404, detail="Project not found")
        ids = [r[0] for r in rows]
        paths = [r[1] for r in rows]
        conn.execute("DELETE FROM local_documents WHERE project = %s", (name,))
    removed = 0
    for rel in paths:
        if _unlink_file_and_prune(rel):
            removed += 1
    # Drop empty project folder under library if present
    proj_dir = _safe_data_path(f"library/{name}")
    if proj_dir and proj_dir.is_dir():
        try:
            for child in sorted(proj_dir.rglob("*"), key=lambda p: len(p.parts), reverse=True):
                if child.is_dir():
                    try:
                        child.rmdir()
                    except OSError:
                        pass
            proj_dir.rmdir()
        except OSError:
            pass
    log.info("Deleted project %s (%d docs, %d files removed)", name, len(ids), removed)
    return {"deleted": True, "project": name, "documentsDeleted": len(ids), "filesRemoved": removed}


@app.post("/documents/{doc_id}/reindex", dependencies=[Depends(require_api_key)])
def reindex(doc_id: str):
    with pool.connection() as conn:
        row = conn.execute(
            "SELECT name, file_path FROM local_documents WHERE id = %s", (doc_id,)
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    src = DATA_DIR / row[1]
    if not src.exists():
        raise HTTPException(status_code=410, detail="File missing on disk")
    with pool.connection() as conn:
        conn.execute(
            "UPDATE local_documents SET status = 'pending', error = NULL, updated_at = now() WHERE id = %s",
            (doc_id,),
        )
    enqueue_embed(doc_id, src, row[0])
    return {"queued": True, "embedQueue": _embed_queue.qsize()}
